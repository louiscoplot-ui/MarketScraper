"""Word .docx letter rendering for the Appraisal Pipeline.

Layout matches the standard Acton | Belle Property Cottesloe template:
  • Full-width brand-green header band with the Acton | belle wordmark
    (auto-uses backend/assets/acton_belle_logo.png if present, else a
    styled text fallback)
  • Right-aligned date
  • Personalised salutation + body
  • Signature block (agent name, role, contact)
  • Page footer in small grey: full agency address, ABN, website
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Acton | Belle Property official brand green — exact value extracted
# from the official wordmark PNG. Was '386351' (off by 1 in the last
# nibble) which read as a slightly different green vs. the logo file.
BRAND_GREEN = '386350'

# Path to the logo. Resolved relative to this module so it works
# regardless of where the Flask process is launched from. Drop the PNG
# at backend/static/logo_acton_belle.png — if missing, the header
# falls through to the styled text wordmark.
_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(_THIS_DIR, 'static', 'logo_acton_belle.png')

# Secondary agency footer lines (street address, contact line, legal/ABN).
# Not exposed in the per-user form yet — env vars override them so an
# operator can swap them without code change. Defaults preserve Louis's
# existing letterhead for backward compatibility.
AGENCY_LINE_1_DEFAULT = '160 Stirling Hwy, Nedlands WA 6009'
AGENCY_LINE_2_DEFAULT = '08 9386 8255  |  cottesloe@belleproperty.com'
AGENCY_LINE_3_DEFAULT = 'Dalkeith Region Pty Ltd  |  ABN 26 123 014 957  |  belleproperty.com/Cottesloe'


def _resolve(profile, profile_key, env_key, default=''):
    """Profile field → env var → default. Profile values that are blank/
    None fall through so a half-completed form doesn't override env vars."""
    if profile:
        v = (profile.get(profile_key) or '').strip()
        if v:
            return v
    return (os.environ.get(env_key) or '').strip() or default


def _format_price(p):
    if p is None or p == '':
        return ''
    try:
        return f"${int(p):,}"
    except (TypeError, ValueError):
        return ''


def _join_oxford(items):
    if len(items) == 0:
        return ''
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ', '.join(items[:-1]) + f", and {items[-1]}"


def format_sources_inline(sources):
    if not sources:
        return '', ''
    n = len(sources)
    addrs = [s['source_address'] for s in sources]
    prices = [_format_price(s.get('source_price')) for s in sources]
    has_prices = [bool(p) for p in prices]
    if n == 1:
        addr_phrase = f"your neighbour at {addrs[0]}"
        sale_phrase = (f"recently sold for {prices[0]}" if has_prices[0]
                       else "recently sold")
    else:
        addr_phrase = f"your neighbours at {_join_oxford(addrs)}"
        if all(has_prices):
            sale_phrase = f"recently sold — for {_join_oxford(prices)} respectively"
        elif any(has_prices):
            paired = [f"{a} ({p})" if p else a for a, p in zip(addrs, prices)]
            addr_phrase = f"your neighbours at {_join_oxford(paired)}"
            sale_phrase = "recently sold"
        else:
            sale_phrase = "recently sold"
    return addr_phrase, sale_phrase


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_padding(cell, top=400, left=720, bottom=400, right=720):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _remove_cell_borders(cell):
    """Strip the default thin borders so the green band reads as a solid block."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _build_text_logo(cell):
    """Fallback wordmark when the logo PNG isn't available — three runs
    in a single paragraph, sized + spaced to mimic the official lockup.
    """
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ACTON — large, light, letter-spaced
    r1 = p.add_run('ACTON')
    r1.font.size = Pt(28)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.name = 'Arial'

    # Spacer + thin separator
    sep = p.add_run('   |   ')
    sep.font.size = Pt(28)
    sep.font.color.rgb = RGBColor(0xC9, 0xD3, 0xCD)  # muted lighter green
    sep.font.name = 'Arial'

    # belle — large, bold, lowercase
    r2 = p.add_run('belle')
    r2.bold = True
    r2.font.size = Pt(28)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2.font.name = 'Arial'

    # PROPERTY — small caps, on the same line for simplicity
    r3 = p.add_run('  PROPERTY')
    r3.bold = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r3.font.name = 'Arial'


def _build_image_logo(cell, logo_path):
    """Drop the official logo PNG inside the green cell, sized to fit."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    # Height fixes vertical band height; width auto-scaled by python-docx
    run.add_picture(logo_path, height=Cm(2.4))


def _emu_to_twips(emu):
    """Word internal: 914400 EMU/inch, 1440 twips/inch."""
    return int(emu / 914400 * 1440)


def _shade_paragraph(p, hex_color):
    """Apply <w:shd> background fill to a paragraph's pPr. This is what
    Word renders as a full-width coloured bar across the page body."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _green_header(doc):
    """Full-width brand-green bar with the Acton | Belle wordmark
    centered inside. The bar is implemented as a single paragraph
    with <w:shd> shading so it spans the body text width (between
    section margins). The logo PNG ships at backend/static/
    logo_acton_belle.png — when missing, falls back to a styled text
    wordmark on the same green background.

    A 2cm-spaced empty paragraph follows the bar to position the body
    correctly for a tri-fold DL envelope window."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, BRAND_GREEN)

    if os.path.exists(LOGO_PATH):
        run = p.add_run()
        # 1.8cm height, width auto-scales from source aspect ratio
        # (924×295 → ~5.64cm wide). add_picture preserves aspect when
        # only one dimension is given.
        run.add_picture(LOGO_PATH, height=Cm(1.8))
    else:
        # Inline text fallback — three runs sized to mimic the lockup.
        r1 = p.add_run('ACTON')
        r1.font.size = Pt(28); r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r1.font.name = 'Arial'
        sep = p.add_run('   |   ')
        sep.font.size = Pt(28); sep.font.color.rgb = RGBColor(0xC9, 0xD3, 0xCD); sep.font.name = 'Arial'
        r2 = p.add_run('belle')
        r2.bold = True; r2.font.size = Pt(28); r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r2.font.name = 'Arial'
        r3 = p.add_run('  PROPERTY')
        r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r3.font.name = 'Arial'

    # 2cm spacer after the bar — positions "Dear [Name]," correctly
    # for a tri-fold DL envelope window. One empty paragraph with
    # space_after = 2cm drives Word's vertical layout.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Cm(2.0)


def _agency_footer(doc, agency_name, line_1, line_2, line_3):
    footer = doc.sections[0].footer
    for p in list(footer.paragraphs):
        p.clear()

    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if agency_name:
        r = p1.add_run(agency_name)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for line in (line_1, line_2, line_3):
        if not line:
            continue
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        p.paragraph_format.space_after = Pt(0)


def render_letter_docx(target_address, owner_name, source_suburb, sources, user_profile=None):
    """Build a prospecting letter. Agency/agent details come from
    `user_profile` (the calling user's saved profile), falling back to
    env vars, then to empty string. Existing callers that don't pass
    `user_profile` keep working — env vars supply the values."""
    raw = (owner_name or '').strip()
    # Reject the legacy "N/A — verify on Landgate" placeholder that still
    # lives in pipeline_tracking rows from before the fix at source.
    if not raw or raw.lower().startswith('n/a'):
        owner = 'Owner'
    else:
        owner = raw
    addr_phrase, sale_phrase = format_sources_inline(sources)
    multi = len(sources) > 1

    profile = user_profile or {}
    agency_name = _resolve(profile, 'agency_name', 'AGENCY_NAME')
    agent_name = _resolve(profile, 'agent_name', 'AGENT_NAME')
    agent_phone = _resolve(profile, 'agent_phone', 'AGENT_PHONE')
    agent_email = _resolve(profile, 'agent_email', 'AGENT_EMAIL')
    # Role line embeds the agency name when known so the signature
    # reads "Sales Agent | <Agency>" — same shape as the original
    # Acton|Belle template, but tenant-aware.
    agent_role = f'Sales Agent | {agency_name}' if agency_name else 'Sales Agent'
    line_1 = (os.environ.get('AGENCY_ADDRESS') or AGENCY_LINE_1_DEFAULT).strip()
    line_2 = (os.environ.get('AGENCY_CONTACT') or AGENCY_LINE_2_DEFAULT).strip()
    line_3 = (os.environ.get('AGENCY_LEGAL') or AGENCY_LINE_3_DEFAULT).strip()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _green_header(doc)
    _agency_footer(doc, agency_name, line_1, line_2, line_3)

    today = datetime.utcnow().strftime('%d/%m/%Y')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(today)
    r.font.size = Pt(11)
    r.font.name = 'Arial'

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f'Dear {owner},')
    r.font.size = Pt(11); r.font.name = 'Arial'

    doc.add_paragraph()

    def body_para(text=None):
        p = doc.add_paragraph()
        if text is not None:
            r = p.add_run(text)
            r.font.size = Pt(11); r.font.name = 'Arial'
        return p

    body_para('I hope this letter finds you well.')

    p = body_para()
    r = p.add_run(f'I wanted to reach out personally — {addr_phrase} {sale_phrase}')
    r.font.size = Pt(11); r.font.name = 'Arial'
    if source_suburb:
        if multi:
            r2 = p.add_run(f", reflecting strong recent results across {source_suburb}.")
        else:
            r2 = p.add_run(f", one of {source_suburb}'s strongest results this season.")
        r2.font.size = Pt(11); r2.font.name = 'Arial'
    else:
        r2 = p.add_run('.')
        r2.font.size = Pt(11); r2.font.name = 'Arial'

    p = body_para()
    if multi:
        intro = ('With this level of activity on your doorstep and buyer demand '
                 'remaining high, this could be the ideal moment to understand '
                 'what your property at ')
    else:
        intro = (f'With buyer demand remaining high across {source_suburb}, this '
                 f'could be the ideal moment to understand what your property at ')
    r1 = p.add_run(intro); r1.font.size = Pt(11); r1.font.name = 'Arial'
    r2 = p.add_run(target_address); r2.bold = True; r2.font.size = Pt(11); r2.font.name = 'Arial'
    r3 = p.add_run(" is truly worth in today's market.")
    r3.font.size = Pt(11); r3.font.name = 'Arial'

    body_para('I would love to offer you a complimentary, no-obligation market '
              'appraisal at a time that suits you — no pressure, just clarity.')

    body_para("Please don't hesitate to reach out.")

    doc.add_paragraph()
    body_para('Kind regards,')
    doc.add_paragraph()

    if agent_name:
        sig = doc.add_paragraph()
        r = sig.add_run(agent_name)
        r.bold = True; r.font.size = Pt(12); r.font.name = 'Arial'

    if agent_role:
        p = doc.add_paragraph()
        r = p.add_run(agent_role)
        r.font.size = Pt(10); r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if agent_phone or agent_email:
        p = doc.add_paragraph()
        if agent_phone:
            rb = p.add_run('M: '); rb.bold = True; rb.font.size = Pt(10); rb.font.name = 'Arial'
            rp = p.add_run(f'{agent_phone}    '); rp.font.size = Pt(10); rp.font.name = 'Arial'
        if agent_email:
            rb = p.add_run('E: '); rb.bold = True; rb.font.size = Pt(10); rb.font.name = 'Arial'
            re_ = p.add_run(agent_email); re_.font.size = Pt(10); re_.font.name = 'Arial'

    return doc
