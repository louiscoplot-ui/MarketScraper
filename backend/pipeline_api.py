"""Appraisal Pipeline routes — endpoints + helper functions."""

import re
import logging
from datetime import datetime, timedelta, date
from io import BytesIO
from flask import request, jsonify, send_file

from database import get_db, normalize_address, USE_POSTGRES

logger = logging.getLogger(__name__)


_ADDR_RE = re.compile(r'^(\d+)([A-Za-z]?)\s+(.+)$')
INSERT_CHUNK = 50


def _source_limit(days):
    return min(max(days * 3, 5), 60)


def _parse_address(addr):
    if not addr:
        return None
    addr = addr.strip()
    if '/' in addr.split()[0]:
        return None
    m = _ADDR_RE.match(addr)
    if not m:
        return None
    try:
        num = int(m.group(1))
    except ValueError:
        return None
    return num, m.group(2), m.group(3).strip()


def _generate_neighbours(addr):
    parsed = _parse_address(addr)
    if not parsed:
        return []
    num, _suffix, street = parsed
    out = []
    for offset in (-2, -1, 1, 2):
        target = num + offset
        if target <= 0:
            continue
        out.append(f"{target} {street}")
    return out


def _hot_vendors_table_exists(conn):
    try:
        if USE_POSTGRES:
            row = conn.execute(
                "SELECT 1 FROM information_schema.tables "
                "WHERE table_schema = 'public' AND table_name = 'hot_vendor_properties' LIMIT 1"
            ).fetchone()
        else:
            row = conn.execute(
                "SELECT 1 FROM sqlite_master "
                "WHERE type = 'table' AND name = 'hot_vendor_properties' LIMIT 1"
            ).fetchone()
        return bool(row)
    except Exception:
        return False


def _match_hot_vendor(conn, target_address):
    try:
        norm = normalize_address(target_address)
        if not norm:
            return None, None
        row = conn.execute(
            "SELECT current_owner, final_score FROM hot_vendor_properties "
            "WHERE normalized_address = ? "
            "ORDER BY final_score DESC LIMIT 1",
            (norm,)
        ).fetchone()
        if not row:
            return None, None
        d = dict(row)
        return d.get('current_owner'), d.get('final_score')
    except Exception:
        return None, None


def _serialize_entries(rows):
    out = []
    for r in rows:
        d = dict(r)
        for k, v in list(d.items()):
            if isinstance(v, (date, datetime)):
                d[k] = v.isoformat()
        out.append(d)
    return out


def _price_to_int(*candidates):
    for c in candidates:
        if c is None or c == '':
            continue
        s = str(c)
        digits = re.sub(r'[^\d]', '', s)
        if digits:
            try:
                return int(digits)
            except ValueError:
                pass
    return None


def _bulk_insert_pipeline(conn, rows):
    if not rows:
        return 0

    inserted = 0
    cols = ('source_address', 'source_suburb', 'source_sold_date',
            'source_price', 'target_address', 'target_owner_name',
            'hot_vendor_score')
    n_cols = len(cols)

    for i in range(0, len(rows), INSERT_CHUNK):
        chunk = rows[i:i + INSERT_CHUNK]
        single = '(' + ', '.join(['?'] * n_cols) + ", 'sent')"
        values_clause = ', '.join([single] * len(chunk))
        sql = (
            f"INSERT INTO pipeline_tracking ({', '.join(cols)}, status) "
            f"VALUES {values_clause} "
            f"ON CONFLICT (target_address, sent_date) DO NOTHING"
        )
        flat_params = [v for row in chunk for v in row]
        cur = conn.execute(sql, flat_params)
        if cur.rowcount and cur.rowcount > 0:
            inserted += cur.rowcount

    return inserted


# ---------------------------------------------------------------------
# Letter rendering helpers
# ---------------------------------------------------------------------

def _format_price(p):
    if p is None or p == '':
        return ''
    try:
        return f"${int(p):,}"
    except (TypeError, ValueError):
        return ''


def _format_sources_inline(sources):
    """Build the human prose listing N source sales."""
    if not sources:
        return '', ''

    n = len(sources)
    addrs = [s['source_address'] for s in sources]
    prices = [_format_price(s.get('source_price')) for s in sources]
    has_prices = [bool(p) for p in prices]

    def join_oxford(items):
        if len(items) == 0:
            return ''
        if len(items) == 1:
            return items[0]
        if len(items) == 2:
            return f"{items[0]} and {items[1]}"
        return ', '.join(items[:-1]) + f", and {items[-1]}"

    if n == 1:
        addr_phrase = f"your neighbour at {addrs[0]}"
        if has_prices[0]:
            sale_phrase = f"recently sold for {prices[0]}"
        else:
            sale_phrase = "recently sold"
    else:
        addr_phrase = f"your neighbours at {join_oxford(addrs)}"
        if all(has_prices):
            sale_phrase = (
                f"recently sold — for {join_oxford(prices)} respectively"
            )
        elif any(has_prices):
            paired = [
                f"{a} ({p})" if p else a
                for a, p in zip(addrs, prices)
            ]
            addr_phrase = f"your neighbours at {join_oxford(paired)}"
            sale_phrase = "recently sold"
        else:
            sale_phrase = "recently sold"

    return addr_phrase, sale_phrase


BRAND_GREEN = '386351'  # Acton | Belle dark green from official letterhead


def _set_cell_shading(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_row_height(row, twips, exact=True):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = OxmlElement('w:trHeight')
    tr_h.set(qn('w:val'), str(twips))
    if exact:
        tr_h.set(qn('w:hRule'), 'exact')
    tr_pr.append(tr_h)


def _set_cell_margins(cell, top_twips=0, bottom_twips=0, left_twips=0, right_twips=0):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for side, val in (('top', top_twips), ('left', left_twips),
                      ('bottom', bottom_twips), ('right', right_twips)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def _add_top_border(paragraph, color='000000', size=6):
    """Add a horizontal rule above the paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), color)
    p_borders.append(top)
    p_pr.append(p_borders)


def _emu_to_twips(emu):
    """Word internal: 914400 EMU per inch, 1440 twips per inch."""
    return int(emu / 914400 * 1440)


def _render_letter_docx(target_address, owner_name, source_suburb, sources):
    """Build an Acton | Belle Property Cottesloe letter matching the official
    letterhead: dark green band with logo at top, body, footer with company
    details and a black horizontal rule above."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    owner = (owner_name or '').strip() or 'Homeowner'
    addr_phrase, sale_phrase = _format_sources_inline(sources)
    multi = len(sources) > 1

    doc = Document()

    section = doc.sections[0]
    # Body margins. Top is generous to clear the header band, bottom leaves
    # room for the company footer block.
    section.top_margin = Cm(5.5)
    section.bottom_margin = Cm(4.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0.8)

    page_width_twips = _emu_to_twips(section.page_width)
    left_margin_twips = _emu_to_twips(section.left_margin)

    # ---------------- HEADER: full-width green band with logo ----------------
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    header_tbl = header.add_table(rows=1, cols=1, width=section.page_width)
    header_tbl.autofit = False

    # Bleed the table out to the page edges (negative left indent + full
    # page width). Without this, the band stops at the body margins.
    tbl_pr = header_tbl._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        header_tbl._element.insert(0, tbl_pr)

    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), str(-left_margin_twips))
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)

    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), str(page_width_twips))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_w)

    cell = header_tbl.rows[0].cells[0]
    _set_cell_shading(cell, BRAND_GREEN)
    _set_row_height(header_tbl.rows[0], 1700)  # ~3.0cm band

    # Vertical-align the logo in the middle of the band
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), 'center')
    tc_pr.append(v_align)

    # Push logo in from the page edge (~3cm from left)
    _set_cell_margins(cell, left_twips=1700, right_twips=1700)

    logo_p = cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(0)

    def add_logo_run(text, size, bold=False):
        r = logo_p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.bold = bold
        return r

    add_logo_run('ACTON', 22, bold=True)
    add_logo_run('   |   ', 22, bold=False)
    add_logo_run('belle', 24, bold=False)
    add_logo_run('  ', 12, bold=False)
    add_logo_run('PROPERTY', 9, bold=True)

    # ---------------- FOOTER: company info + black rule ----------------
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    rule_p = footer.add_paragraph()
    _add_top_border(rule_p, color='000000', size=6)
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after = Pt(4)

    footer_lines = [
        ('ACTON | Belle Property Cottesloe', 9, True, None),
        ('180 Stirling Hwy, Nedlands WA 6009', 9, False, None),
        ('08 9388 8255', 9, False, None),
        ('cottesloe@belleproperty.com', 9, False, None),
        ('Dalkeith Region Pty Ltd', 9, False, None),
        ('ABN 26 125 014 997', 9, False, None),
        ('belleproperty.com/Cottesloe', 8, False, (0x80, 0x80, 0x80)),
    ]
    for text, size, bold, color in footer_lines:
        fp = footer.add_paragraph(text)
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        for run in fp.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)

    # ---------------- BODY ----------------
    def body_paragraph(text='', size=11, bold=False):
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.font.bold = bold
        return p

    today = datetime.utcnow().strftime('%d %B %Y')
    body_paragraph(today)
    body_paragraph()
    body_paragraph(f'Dear {owner},')
    body_paragraph()
    body_paragraph('I hope this letter finds you well.')

    p = doc.add_paragraph()
    r = p.add_run(f'I wanted to reach out personally — {addr_phrase} {sale_phrase}')
    r.font.name = 'Arial'; r.font.size = Pt(11)
    if source_suburb:
        tail = (f", reflecting strong recent results across {source_suburb}." if multi
                else f", one of {source_suburb}'s strongest results this season.")
    else:
        tail = '.'
    r2 = p.add_run(tail)
    r2.font.name = 'Arial'; r2.font.size = Pt(11)

    p = doc.add_paragraph()
    intro = ('With this level of activity on your doorstep and buyer demand '
             'remaining high, this could be the ideal moment to understand '
             'what your property at ') if multi else \
            (f'With buyer demand remaining high across {source_suburb}, this '
             f'could be the ideal moment to understand what your property at ')
    r = p.add_run(intro); r.font.name = 'Arial'; r.font.size = Pt(11)
    rb = p.add_run(target_address); rb.bold = True; rb.font.name = 'Arial'; rb.font.size = Pt(11)
    re_ = p.add_run(" is truly worth in today's market.")
    re_.font.name = 'Arial'; re_.font.size = Pt(11)

    body_paragraph(
        'I would love to offer you a complimentary, no-obligation market '
        'appraisal at a time that suits you — no pressure, just clarity.'
    )
    body_paragraph("Please don't hesitate to reach out.")
    body_paragraph()
    body_paragraph('Warm regards,')
    body_paragraph()

    sig = doc.add_paragraph()
    rs = sig.add_run('Louis Coplot')
    rs.bold = True; rs.font.name = 'Arial'; rs.font.size = Pt(13)

    body_paragraph('Sales Agent | Belle Property Cottesloe', size=10)
    body_paragraph('M: 0400 XXX XXX', size=10)
    body_paragraph('E: louis@belleproperty.com.au', size=10)
    body_paragraph('W: belleproperty.com.au/cottesloe', size=10)

    return doc


def _gather_sources_for_target(conn, target_address, source_suburb):
    """Distinct nearby sales for a target. Dedupes by source_address so
    duplicate pipeline_tracking rows (different sent_dates, same source)
    don't make the letter say 'sold for X — for X respectively'."""
    rows = conn.execute(
        """
        SELECT source_address, source_price, source_sold_date, target_owner_name
        FROM pipeline_tracking
        WHERE LOWER(target_address) = LOWER(?)
          AND LOWER(source_suburb) = LOWER(?)
        ORDER BY created_at DESC
        """,
        (target_address, source_suburb)
    ).fetchall()

    seen = set()
    deduped = []
    for r in rows:
        d = dict(r)
        key = (d.get('source_address') or '').strip().lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(d)
    deduped.sort(key=lambda d: (d.get('source_address') or '').lower())
    return deduped


# ---------------------------------------------------------------------
# Route handlers
# ---------------------------------------------------------------------

def pipeline_generate():
    suburb = (request.args.get('suburb') or '').strip()
    if not suburb:
        return jsonify({'error': 'suburb is required'}), 400

    try:
        days = int(request.args.get('days') or 7)
    except ValueError:
        days = 7
    days = max(1, min(days, 90))

    cutoff_date = (datetime.utcnow() - timedelta(days=days)).date().isoformat()
    src_limit = _source_limit(days)

    conn = get_db()

    has_hv = _hot_vendors_table_exists(conn)

    sold_rows = conn.execute(
        """
        SELECT l.address, l.sold_price, l.price_text, l.sold_date,
               l.first_seen, l.last_seen, s.name AS suburb_name,
               COALESCE(l.sold_date, SUBSTR(l.first_seen, 1, 10)) AS effective_date
        FROM listings l
        JOIN suburbs s ON l.suburb_id = s.id
        WHERE l.status = 'sold'
          AND LOWER(s.name) = LOWER(?)
          AND COALESCE(l.sold_date, SUBSTR(l.first_seen, 1, 10)) >= ?
        ORDER BY COALESCE(l.sold_date, SUBSTR(l.first_seen, 1, 10)) DESC,
                 l.first_seen DESC
        LIMIT ?
        """,
        (suburb, cutoff_date, src_limit)
    ).fetchall()

    sold_count = len(sold_rows)

    insert_rows = []
    for r in sold_rows:
        source_address = r['address']
        source_suburb = r['suburb_name']
        source_price = _price_to_int(r['sold_price'], r['price_text'])
        source_sold_date = r['effective_date']

        for target in _generate_neighbours(source_address):
            if has_hv:
                owner, score = _match_hot_vendor(conn, target)
            else:
                owner, score = None, None
            insert_rows.append((
                source_address, source_suburb, source_sold_date,
                source_price, target, owner, score,
            ))

    generated = _bulk_insert_pipeline(conn, insert_rows)
    conn.commit()

    rows = conn.execute(
        """
        SELECT * FROM pipeline_tracking
        WHERE LOWER(source_suburb) = LOWER(?)
        ORDER BY created_at DESC
        LIMIT 500
        """,
        (suburb,)
    ).fetchall()
    entries = _serialize_entries(rows)
    conn.close()

    return jsonify({
        'generated': generated,
        'sold_count': sold_count,
        'suburb': suburb,
        'cap_applied': sold_count >= src_limit,
        'entries': entries,
    })


def pipeline_manual_add():
    """POST /api/pipeline/manual-add — create pipeline entries from a sale
    the agent knows about that the scraper either missed or hasn't dated
    correctly yet."""
    data = request.get_json(silent=True) or {}

    source_address = (data.get('source_address') or '').strip()
    source_suburb = (data.get('source_suburb') or '').strip()
    if not source_address:
        return jsonify({'error': 'source_address is required'}), 400
    if not source_suburb:
        return jsonify({'error': 'source_suburb is required'}), 400

    source_price = _price_to_int(data.get('source_price'))
    source_sold_date = (data.get('source_sold_date') or '').strip() or \
        date.today().isoformat()

    explicit_targets = data.get('target_addresses') or []
    if not isinstance(explicit_targets, list):
        return jsonify({'error': 'target_addresses must be a list'}), 400
    explicit_targets = [
        t.strip() for t in explicit_targets
        if isinstance(t, str) and t.strip()
    ]

    if explicit_targets:
        targets = explicit_targets
    else:
        targets = _generate_neighbours(source_address)
        if not targets:
            return jsonify({
                'error': (
                    f"Couldn't auto-generate neighbours from '{source_address}' "
                    "(strata or non-numeric address). Provide target_addresses "
                    "explicitly."
                )
            }), 400

    conn = get_db()
    has_hv = _hot_vendors_table_exists(conn)

    insert_rows = []
    for target in targets:
        if has_hv:
            owner, score = _match_hot_vendor(conn, target)
        else:
            owner, score = None, None
        insert_rows.append((
            source_address, source_suburb, source_sold_date,
            source_price, target, owner, score,
        ))

    generated = _bulk_insert_pipeline(conn, insert_rows)
    conn.commit()

    rows = conn.execute(
        """
        SELECT * FROM pipeline_tracking
        WHERE LOWER(source_address) = LOWER(?)
          AND LOWER(source_suburb) = LOWER(?)
          AND sent_date = ?
        ORDER BY target_address ASC
        """,
        (source_address, source_suburb, date.today().isoformat())
    ).fetchall()
    entries = _serialize_entries(rows)
    conn.close()

    return jsonify({
        'generated': generated,
        'attempted': len(targets),
        'targets': targets,
        'source_address': source_address,
        'source_suburb': source_suburb,
        'entries': entries,
    })


def pipeline_tracking_list():
    suburb = (request.args.get('suburb') or '').strip()
    status = (request.args.get('status') or '').strip()
    try:
        limit = int(request.args.get('limit') or 100)
    except ValueError:
        limit = 100
    limit = max(1, min(limit, 1000))

    conn = get_db()
    sql = "SELECT * FROM pipeline_tracking WHERE 1=1"
    params = []
    if suburb:
        sql += " AND LOWER(source_suburb) = LOWER(?)"
        params.append(suburb)
    if status:
        sql += " AND status = ?"
        params.append(status)
    sql += " ORDER BY created_at DESC LIMIT ?"
    params.append(limit)

    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify({'entries': _serialize_entries(rows)})


def pipeline_tracking_grouped():
    """One envelope per target_address. Sources are deduped (same source
    appearing across multiple sent_dates collapses to one entry) and
    sorted alphabetically inside each group. Groups themselves are
    ordered by their primary source's address — so the user can scan
    the pipeline by the SALES they're targeting, not by neighbour."""
    suburb = (request.args.get('suburb') or '').strip()
    try:
        limit = int(request.args.get('limit') or 200)
    except ValueError:
        limit = 200
    limit = max(1, min(limit, 1000))

    conn = get_db()
    sql = "SELECT * FROM pipeline_tracking WHERE 1=1"
    params = []
    if suburb:
        sql += " AND LOWER(source_suburb) = LOWER(?)"
        params.append(suburb)
    sql += " ORDER BY target_address ASC, created_at DESC"

    rows = conn.execute(sql, params).fetchall()
    conn.close()

    groups = {}
    for raw in rows:
        r = dict(raw)
        for k, v in list(r.items()):
            if isinstance(v, (date, datetime)):
                r[k] = v.isoformat()
        key = (r['target_address'].lower().strip(), (r.get('source_suburb') or '').lower().strip())
        if key not in groups:
            groups[key] = {
                'target_address': r['target_address'],
                'source_suburb': r.get('source_suburb'),
                'target_owner_name': r.get('target_owner_name'),
                'hot_vendor_score': r.get('hot_vendor_score'),
                'status': r.get('status'),
                'sent_date': r.get('sent_date'),
                'response_date': r.get('response_date'),
                'notes': r.get('notes'),
                'representative_id': r.get('id'),
                'row_ids': [],
                'sources': [],
                '_source_addrs_seen': set(),
            }
        g = groups[key]
        g['row_ids'].append(r['id'])
        src_addr_key = (r.get('source_address') or '').strip().lower()
        if src_addr_key and src_addr_key not in g['_source_addrs_seen']:
            g['_source_addrs_seen'].add(src_addr_key)
            g['sources'].append({
                'source_address': r.get('source_address'),
                'source_price': r.get('source_price'),
                'source_sold_date': r.get('source_sold_date'),
                'row_id': r.get('id'),
            })
        if not g.get('target_owner_name') and r.get('target_owner_name'):
            g['target_owner_name'] = r.get('target_owner_name')
        if g.get('hot_vendor_score') is None and r.get('hot_vendor_score') is not None:
            g['hot_vendor_score'] = r.get('hot_vendor_score')

    grouped = list(groups.values())
    # Sort sources within each group alphabetically
    for g in grouped:
        g['sources'].sort(key=lambda s: (s.get('source_address') or '').lower())
        g.pop('_source_addrs_seen', None)
    # Sort groups by primary source's address (the one we're targeting via
    # this letter). Falls back to target_address when sources is empty.
    grouped.sort(key=lambda g: (
        (g['sources'][0].get('source_address') if g['sources'] else g['target_address'] or '').lower()
    ))
    grouped = grouped[:limit]
    return jsonify({'groups': grouped, 'count': len(grouped)})


def pipeline_tracking_update(id):
    data = request.get_json(silent=True) or {}
    allowed = ('status', 'response_date', 'notes', 'target_owner_name')

    sets = []
    params = []
    for key in allowed:
        if key in data:
            sets.append(f"{key} = ?")
            params.append(data[key])

    if not sets:
        return jsonify({'error': 'No updatable fields provided'}), 400

    propagate_owner = 'target_owner_name' in data

    conn = get_db()

    target_row = conn.execute(
        "SELECT target_address, source_suburb FROM pipeline_tracking WHERE id = ?",
        (id,)
    ).fetchone()
    if not target_row:
        conn.close()
        return jsonify({'error': 'Not found'}), 404

    if propagate_owner:
        conn.execute(
            "UPDATE pipeline_tracking SET target_owner_name = ? "
            "WHERE LOWER(target_address) = LOWER(?) AND LOWER(source_suburb) = LOWER(?)",
            (data['target_owner_name'], target_row['target_address'], target_row['source_suburb'])
        )

    other_sets = [s for s in sets if not s.startswith('target_owner_name')]
    if other_sets:
        other_params = []
        for key in allowed:
            if key == 'target_owner_name':
                continue
            if key in data:
                other_params.append(data[key])
        other_params.append(id)
        conn.execute(
            f"UPDATE pipeline_tracking SET {', '.join(other_sets)} WHERE id = ?",
            other_params
        )

    conn.commit()

    row = conn.execute(
        "SELECT * FROM pipeline_tracking WHERE id = ?", (id,)
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({'error': 'Not found'}), 404

    return jsonify(_serialize_entries([row])[0])


def pipeline_tracking_clear():
    if (request.args.get('confirm') or '').lower() != 'yes':
        return jsonify({
            'error': 'Add ?confirm=yes to actually delete. This is destructive.'
        }), 400

    suburb = (request.args.get('suburb') or '').strip()
    status = (request.args.get('status') or '').strip()

    if not suburb and not status:
        return jsonify({
            'error': 'At least suburb or status must be provided. '
                     'Refusing to wipe the entire table.'
        }), 400

    conn = get_db()
    sql = "DELETE FROM pipeline_tracking WHERE 1=1"
    params = []
    if suburb:
        sql += " AND LOWER(source_suburb) = LOWER(?)"
        params.append(suburb)
    if status:
        sql += " AND status = ?"
        params.append(status)

    cur = conn.execute(sql, params)
    deleted = cur.rowcount or 0
    conn.commit()
    conn.close()

    return jsonify({
        'deleted': deleted,
        'suburb': suburb or None,
        'status': status or None,
    })


def pipeline_letter_download(id):
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM pipeline_tracking WHERE id = ?", (id,)
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({'error': 'Letter not found'}), 404

    entry = dict(row)
    target_address = entry.get('target_address') or ''
    source_suburb = entry.get('source_suburb') or ''

    sources = _gather_sources_for_target(conn, target_address, source_suburb)
    conn.close()

    if not sources:
        sources = [{
            'source_address': entry.get('source_address'),
            'source_price': entry.get('source_price'),
            'source_sold_date': entry.get('source_sold_date'),
        }]

    owner_name = entry.get('target_owner_name')
    if not owner_name:
        for s in sources:
            n = (s.get('target_owner_name') or '').strip()
            if n:
                owner_name = n
                break

    doc = _render_letter_docx(target_address, owner_name, source_suburb, sources)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe = re.sub(r'[^\w\s-]', '', target_address)[:60].strip().replace(' ', '_')
    filename = f"letter_{safe or 'letter'}.docx"

    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename,
    )


# ---------------------------------------------------------------------
# Wiring
# ---------------------------------------------------------------------

def register_pipeline_routes(app):
    app.add_url_rule('/api/pipeline/generate', endpoint='pipeline_generate',
                     view_func=pipeline_generate, methods=['GET'])
    app.add_url_rule('/api/pipeline/manual-add', endpoint='pipeline_manual_add',
                     view_func=pipeline_manual_add, methods=['POST'])
    app.add_url_rule('/api/pipeline/tracking', endpoint='pipeline_tracking_list',
                     view_func=pipeline_tracking_list, methods=['GET'])
    app.add_url_rule('/api/pipeline/tracking/grouped', endpoint='pipeline_tracking_grouped',
                     view_func=pipeline_tracking_grouped, methods=['GET'])
    app.add_url_rule('/api/pipeline/tracking/<int:id>', endpoint='pipeline_tracking_update',
                     view_func=pipeline_tracking_update, methods=['PATCH'])
    app.add_url_rule('/api/pipeline/tracking', endpoint='pipeline_tracking_clear',
                     view_func=pipeline_tracking_clear, methods=['DELETE'])
    app.add_url_rule('/api/pipeline/letter/<int:id>/download',
                     endpoint='pipeline_letter_download',
                     view_func=pipeline_letter_download, methods=['GET'])
    logger.info(
        "Pipeline routes: /api/pipeline/{generate,manual-add,tracking[,/grouped,/<id>],letter/<id>/download}"
    )
