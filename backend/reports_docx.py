"""Reports .docx rendering — the 4 branded intelligence reports.

Reuses the Acton | Belle letter chrome from pipeline_letter.py verbatim
(full-bleed BRAND_GREEN #386350 header band with the logo, agency footer
block, Arial body, A4 portrait) so a report and a letter printed the
same morning look like one document family.

Design goals (this is a client-facing deliverable, not a data dump):
  * Colour-coded KPI cards, temperature gauges and horizontal bars —
    all drawn with native docx primitives (shaded cells + coloured
    block-glyph runs), so NO new dependency (matplotlib etc.).
  * Plain-English "What this means" observations derived deterministically
    from the numbers, so the report explains itself even when the AI
    narrative layer is unavailable (no ANTHROPIC_API_KEY).
  * One page per suburb; technical caveats collapsed into a single small
    footnote instead of a wall of warnings.

Data-honesty rules (enforced upstream in reports_engine, respected here):
  * A None metric renders as "n/a" — never a substituted number.
  * The vendor_benchmark variant is vendor-safe: no agency names, no
    competitor addresses.
  * The rentals block in the deep dive is an explicit placeholder.
"""

import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from time_utils import perth_now
from pipeline_letter import (
    BRAND_GREEN, _green_header, _agency_footer, _resolve, _shade_cell,
    _remove_cell_borders, _set_cell_padding,
    AGENCY_LINE_1_DEFAULT, AGENCY_LINE_2_DEFAULT, AGENCY_LINE_3_DEFAULT,
)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Palette — brand green + the status grammar mirrored from the frontend
# STATUS_COLORS (App.jsx:31): good / watch / info / alert.
GREEN = RGBColor(0x38, 0x63, 0x50)
GREY = RGBColor(0x55, 0x55, 0x55)
GREY_LIGHT = RGBColor(0x8A, 0x93, 0x8C)
INK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BAR_EMPTY = RGBColor(0xE5, 0xE7, 0xEB)
HOT = RGBColor(0x15, 0x80, 0x3D)     # seller's market
WARM = RGBColor(0xB4, 0x53, 0x09)    # balanced
COOL = RGBColor(0xB9, 0x1C, 0x1C)    # buyer's market
BLUE = RGBColor(0x25, 0x63, 0xEB)

# Light card fills.
FILL_CARD = 'F4F6F5'
FILL_HOT = 'DCFCE7'
FILL_WARM = 'FEF3C7'
FILL_COOL = 'FEE2E2'
FILL_NEUTRAL = 'F1F3F5'

# Stacked stock-aging bar colours (fresh → stale).
AGE_COLORS = [RGBColor(0x16, 0xA3, 0x4A), RGBColor(0x84, 0xCC, 0x16),
              RGBColor(0xD9, 0x77, 0x06), RGBColor(0xDC, 0x26, 0x26)]


# ---------------------------------------------------------------- low-level

def _set_cell_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcW'))
    if old is not None:
        tcPr.remove(old)
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(twips))
    w.set(qn('w:type'), 'dxa')
    tcPr.append(w)


def _cell_v_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)


def _run(p, text, size=10.5, color=INK, bold=False, italic=False, name='Arial'):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = name
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _para(doc, space_after=4, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    return p


# ---------------------------------------------------------------- chrome

def _new_report_doc(user_profile, title, subtitle):
    profile = user_profile or {}
    agency_name = _resolve(profile, 'agency_name', 'AGENCY_NAME')
    agent_name = _resolve(profile, 'agent_name', 'AGENT_NAME')
    line_1 = (os.environ.get('AGENCY_ADDRESS') or AGENCY_LINE_1_DEFAULT).strip()
    line_2 = (os.environ.get('AGENCY_CONTACT') or AGENCY_LINE_2_DEFAULT).strip()
    line_3 = (os.environ.get('AGENCY_LEGAL') or AGENCY_LINE_3_DEFAULT).strip()

    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    _green_header(doc)
    _agency_footer(doc, agency_name, line_1, line_2, line_3)
    # Trim the letter's 2cm envelope-window spacer.
    doc.paragraphs[-1].paragraph_format.space_after = Cm(0.25)

    p = _para(doc, space_after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _run(p, perth_now().strftime('%d %B %Y'), size=9.5, color=GREY_LIGHT)
    tp = _para(doc, space_after=1)
    _run(tp, title, size=21, color=GREEN, bold=True)
    sp = _para(doc, space_after=2)
    _run(sp, subtitle, size=10, color=GREY)
    if agent_name:
        pp = _para(doc, space_after=6)
        _run(pp, 'Prepared by ' + agent_name
             + (' · ' + agency_name if agency_name else ''),
             size=9, color=GREY_LIGHT)
    _hrule(doc)
    return doc


def _hrule(doc, color='D8DEDB'):
    p = _para(doc, space_after=6, space_before=0)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pbdr.append(b)
    pPr.append(pbdr)


def _section(doc, text):
    p = _para(doc, space_before=10, space_after=4)
    _run(p, text.upper(), size=10, color=GREEN, bold=True)


# ---------------------------------------------------------------- widgets

def _kpi_cards(doc, cards):
    """A row of colour-coded KPI cards. `cards`: list of
    {value, label, sub, color, fill}. Thin spacer columns separate them."""
    n = len(cards)
    cols = n * 2 - 1
    t = doc.add_table(rows=1, cols=cols)
    t.autofit = False
    # Content width ≈ 17cm → 9639 twips. Cards share it, gaps ~160tw.
    gap = 160
    card_w = int((9639 - gap * (n - 1)) / n)
    ci = 0
    for idx, c in enumerate(cards):
        cell = t.cell(0, ci)
        _remove_cell_borders(cell)
        _shade_cell(cell, c.get('fill', FILL_CARD))
        _set_cell_padding(cell, top=150, left=190, bottom=150, right=120)
        _set_cell_width(cell, card_w)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, c['value'], size=21, color=c.get('color', GREEN), bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(0)
        _run(p2, c['label'].upper(), size=7.5, color=GREY)
        if c.get('sub'):
            p3 = cell.add_paragraph()
            p3.paragraph_format.space_before = Pt(1)
            p3.paragraph_format.space_after = Pt(0)
            _run(p3, c['sub'], size=8, color=c.get('color', GREY), bold=True)
        ci += 1
        if idx < n - 1:
            g = t.cell(0, ci)
            _remove_cell_borders(g)
            _set_cell_width(g, gap)
            ci += 1


def _bar_rows(doc, rows, label_w=3.4, bar_blocks=26):
    """Borderless 3-col table of horizontal bars: label | block-bar | value.
    `rows`: list of {label, pct, color, value}."""
    t = doc.add_table(rows=0, cols=3)
    t.autofit = False
    lw = int(label_w * 567)
    vw = int(2.6 * 567)
    bw = 9639 - lw - vw
    for row in rows:
        cells = t.add_row().cells
        for cc in cells:
            _remove_cell_borders(cc)
            _cell_v_center(cc)
        _set_cell_width(cells[0], lw)
        _set_cell_width(cells[1], bw)
        _set_cell_width(cells[2], vw)
        pl = cells[0].paragraphs[0]
        pl.paragraph_format.space_after = Pt(1)
        _run(pl, row['label'], size=9, color=INK)
        pb = cells[1].paragraphs[0]
        pb.paragraph_format.space_after = Pt(1)
        pct = max(0.0, min(100.0, row.get('pct') or 0))
        filled = round(pct / 100 * bar_blocks)
        if filled:
            _run(pb, '█' * filled, size=9, color=row.get('color', GREEN))
        if bar_blocks - filled:
            _run(pb, '█' * (bar_blocks - filled), size=9, color=BAR_EMPTY)
        pv = cells[2].paragraphs[0]
        pv.paragraph_format.space_after = Pt(1)
        pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(pv, str(row['value']), size=9, color=INK, bold=True)


def _stacked_bar(doc, segments, total_blocks=40):
    """Single-line stacked bar from coloured segments.
    `segments`: list of (count, color). Zero-total → light placeholder."""
    total = sum(s[0] for s in segments)
    p = _para(doc, space_after=2)
    if total <= 0:
        _run(p, '█' * total_blocks, size=10, color=BAR_EMPTY)
        return
    used = 0
    for i, (count, color) in enumerate(segments):
        if count <= 0:
            continue
        blocks = round(count / total * total_blocks)
        if i == len(segments) - 1:
            blocks = total_blocks - used  # absorb rounding into the last
        blocks = max(0, blocks)
        used += blocks
        if blocks:
            _run(p, '█' * blocks, size=10, color=color)


def _gauge(doc, score):
    """Momentum temperature gauge — a wide block bar coloured by band."""
    color, _fill, _word = _temp(score)
    p = _para(doc, space_after=2)
    blocks = 46
    pct = 0 if score is None else max(0, min(100, score))
    filled = round(pct / 100 * blocks)
    if filled:
        _run(p, '█' * filled, size=11, color=color)
    if blocks - filled:
        _run(p, '█' * (blocks - filled), size=11, color=BAR_EMPTY)
    cap = _para(doc, space_after=4)
    _run(cap, '  Cooler / buyer’s market', size=7.5, color=GREY_LIGHT)
    _run(cap, '            Balanced', size=7.5, color=GREY_LIGHT)
    _run(cap, '            Hotter / seller’s market', size=7.5, color=GREY_LIGHT)


def _table(doc, headers, rows, widths=None):
    """Clean bordered table with a brand-green header row."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        _shade_cell(hdr[i], BRAND_GREEN)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, str(h), size=8.5, color=WHITE, bold=True)
        if widths:
            _set_cell_width(hdr[i], int(widths[i] * 567))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            if ri % 2 == 1:
                _shade_cell(cells[i], 'F7F9F8')
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            txt, col, bold = v if isinstance(v, tuple) else (v, INK, False)
            _run(p, 'n/a' if txt is None else str(txt), size=8.5,
                 color=col, bold=bold)
            if widths:
                _set_cell_width(cells[i], int(widths[i] * 567))
    return t


def _observations(doc, lines):
    if not lines:
        return
    _section(doc, 'What this means')
    for ln in lines:
        p = _para(doc, space_after=3)
        _run(p, '›  ', size=10, color=GREEN, bold=True)
        _run(p, ln, size=10, color=INK)


def _notes(doc, flags):
    """Collapse every caveat into ONE small footnote, de-jargoned."""
    if not flags:
        return
    seen, human = set(), []
    for f in flags:
        key = f.lower()
        if 'sample too small' in key or 'low ' in key or 'no data' in key \
                or 'withheld' in key or 'indicative' in key:
            msg = 'some figures rest on small samples and are marked n/a or indicative'
        elif 'baseline' in key:
            msg = 'momentum baseline is partial where less than 12 months of history exists'
        elif 'not captured' in key or 'unavailable' in key or 'not tracked' in key:
            msg = 'a few metrics are shown as n/a where the source data was not captured'
        else:
            continue
        if msg not in seen:
            seen.add(msg)
            human.append(msg)
    if not human:
        return
    p = _para(doc, space_before=6, space_after=0)
    _run(p, 'Notes: ' + '; '.join(human) + '.', size=7.5,
         color=GREY_LIGHT, italic=True)


def _narrative(doc, narratives, key):
    text = (narratives or {}).get(key)
    if text:
        p = _para(doc, space_before=3, space_after=6)
        _run(p, text, size=10, color=INK, italic=True)


# ---------------------------------------------------------------- helpers

def _temp(score):
    """Momentum band → (text colour, light fill, word)."""
    if score is None:
        return GREY_LIGHT, FILL_NEUTRAL, 'No score'
    if score >= 67:
        return HOT, FILL_HOT, 'Hotter'
    if score >= 34:
        return WARM, FILL_WARM, 'Balanced'
    return COOL, FILL_COOL, 'Cooler'


def _fmt(v, suffix=''):
    return 'n/a' if v is None else f'{v}{suffix}'


def _collect_flags(m):
    out = []
    for block in ('momentum', 'months_of_supply', 'velocity', 'discount',
                  'agency_share', 'stock_aging', 'stale_flags'):
        b = m.get(block) or {}
        out.extend(b.get('flags', []) or [])
    for band in m.get('price_bands', []) or []:
        out.extend(band.get('flags', []) or [])
    return out


def _suburb_observations(m):
    """Deterministic plain-English read of the numbers — grounded strictly
    in computed values, never invents. Renders even without the AI layer."""
    obs = []
    mom = m['momentum']
    sc = mom.get('score')
    _c, _f, word = _temp(sc)
    if sc is not None:
        if sc >= 67:
            obs.append(f"Momentum is running hotter than usual at {sc}/100 — "
                       "conditions currently favour sellers versus this "
                       "suburb’s own 12-month norm.")
        elif sc >= 34:
            obs.append(f"Momentum sits balanced at {sc}/100, broadly in line "
                       "with this suburb’s 12-month norm.")
        else:
            obs.append(f"Momentum is cooler at {sc}/100 — buyers currently "
                       "hold more leverage than this suburb’s 12-month norm.")
    mos = m['months_of_supply'].get('value')
    if mos is not None:
        if mos >= 6:
            obs.append(f"At {mos} months of supply there is a comfortable "
                       "runway of stock, so pricing and presentation carry "
                       "more weight to secure a sale.")
        elif mos <= 3:
            obs.append(f"At {mos} months of supply stock is tight, and "
                       "well-presented homes can move quickly.")
        else:
            obs.append(f"At {mos} months of supply the balance between buyers "
                       "and sellers is fairly even.")
    dom = m['velocity'].get('dom_median_active')
    if dom is not None:
        obs.append(f"Homes currently for sale have a median of {dom} days on "
                   "market.")
    sa = m['stock_aging']
    tot = sa.get('total_active_aged') or 0
    if tot:
        over60 = sa['buckets']['d60_90'] + sa['buckets']['over_90d']
        pct = round(over60 / tot * 100)
        if pct >= 20:
            obs.append(f"{pct}% of homes for sale have been listed more than "
                       "60 days — a pool of potentially motivated vendors.")
    d = m['discount']
    if d.get('available'):
        obs.append(f"Recent sales settled a median {d['median_pct']}% below "
                   f"their original asking price (last {d['window_days']} days).")
    sf = m['stale_flags']
    if sf.get('count'):
        obs.append(f"{sf['count']} active campaign(s) are both long-running "
                   "and have already reduced price — worth a direct approach.")
    return obs[:6]


# ------------------------------------------------------------ per-suburb

def _suburb_page(doc, m, narratives, first):
    if not first:
        doc.add_page_break()
    if m.get('error'):
        _section(doc, m.get('suburb') or f"Suburb #{m.get('suburb_id')}")
        p = _para(doc)
        _run(p, 'Data temporarily unavailable for this suburb.', italic=True,
             color=GREY)
        return

    sc = m['momentum'].get('score')
    tcol, tfill, tword = _temp(sc)
    c = m['counts']

    # Suburb headline
    hp = _para(doc, space_after=1)
    _run(hp, m.get('suburb') or f"Suburb #{m.get('suburb_id')}", size=17,
         color=GREEN, bold=True)
    sub = _para(doc, space_after=6)
    _run(sub, f"{c['active']} for sale  ·  {c['under_offer']} under offer  ·  "
              f"{c['sold']} sold on record  ·  {c['withdrawn']} withdrawn",
         size=9, color=GREY_LIGHT)

    # KPI cards
    dom = m['velocity'].get('dom_median_active')
    mos = m['months_of_supply'].get('value')
    disc = m['discount'].get('median_pct') if m['discount'].get('available') else None
    _kpi_cards(doc, [
        {'value': _fmt(sc, '/100'), 'label': 'Momentum', 'sub': tword,
         'color': tcol, 'fill': tfill},
        {'value': _fmt(dom, 'd'), 'label': 'Median days on market',
         'color': INK, 'fill': FILL_CARD},
        {'value': _fmt(mos), 'label': 'Months of supply',
         'color': INK, 'fill': FILL_CARD},
        {'value': _fmt(disc, '%') if disc is not None else 'n/a',
         'label': 'List-to-sale discount', 'color': INK, 'fill': FILL_CARD},
    ])

    # Momentum gauge
    _section(doc, 'Market momentum')
    _gauge(doc, sc)

    # Stock aging stacked bar
    b = m['stock_aging']['buckets']
    _section(doc, 'Stock on market by age')
    _stacked_bar(doc, [
        (b['under_30d'], AGE_COLORS[0]),
        (b['d30_60'], AGE_COLORS[1]),
        (b['d60_90'], AGE_COLORS[2]),
        (b['over_90d'], AGE_COLORS[3]),
    ])
    lp = _para(doc, space_after=4)
    _run(lp, f"■ <30 days ({b['under_30d']})   ", size=7.5, color=AGE_COLORS[0])
    _run(lp, f"■ 30–60 ({b['d30_60']})   ", size=7.5, color=AGE_COLORS[1])
    _run(lp, f"■ 60–90 ({b['d60_90']})   ", size=7.5, color=AGE_COLORS[2])
    _run(lp, f"■ 90+ ({b['over_90d']})", size=7.5, color=AGE_COLORS[3])

    # Agency share (agency level only) — horizontal bars, top 5
    ash = m['agency_share']
    roll = ash['rolling_90d'][:5]
    if roll:
        _section(doc, 'Share of new listings · rolling 3 months (agency)')
        maxpct = max((x['pct'] for x in roll), default=1) or 1
        _bar_rows(doc, [{
            'label': x['agency'][:34],
            'pct': x['pct'] / maxpct * 100,
            'color': GREEN,
            'value': f"{x['count']} · {x['pct']}%",
        } for x in roll])

    # Observations — AI narrative if present, else deterministic read
    ai = (narratives or {}).get('momentum') or (narratives or {}).get('house_view')
    if ai:
        _section(doc, 'What this means')
        _narrative(doc, narratives, 'momentum')
        _narrative(doc, narratives, 'pricing')
    else:
        _observations(doc, _suburb_observations(m))

    _notes(doc, _collect_flags(m))


# ------------------------------------------------------------ generators

def build_suburb_intelligence(metrics, narratives, user_profile):
    doc = _new_report_doc(
        user_profile, 'Suburb Intelligence',
        'Momentum, velocity, pricing, stock and competitive position — '
        'one page per suburb.')
    for i, m in enumerate(metrics):
        _suburb_page(doc, m, narratives, first=(i == 0))
    return doc


def build_director_dashboard(metrics, narratives, user_profile):
    doc = _new_report_doc(
        user_profile, 'Director Dashboard',
        'Portfolio momentum, listing-share movement and the opportunity '
        'register across your suburbs.')
    ok = [m for m in metrics if not m.get('error')]

    _section(doc, 'Momentum heat map')
    rows = []
    for m in sorted(ok, key=lambda x: -(x['momentum'].get('score') or -1)):
        sc = m['momentum'].get('score')
        tcol, _f, word = _temp(sc)
        rows.append([
            (m['suburb'], INK, True),
            (_fmt(sc, '/100'), tcol, True),
            (word, tcol, False),
            _fmt(m['months_of_supply'].get('value')),
            _fmt(m['velocity'].get('dom_median_active'), 'd'),
            (f"{m['velocity']['recent_cohort'].get('pct_absorbed_21d')}%"
             if m['velocity']['recent_cohort'].get('pct_absorbed_21d') is not None
             else 'n/a'),
            m['counts']['active'],
        ])
    _table(doc, ['Suburb', 'Momentum', 'Temperature', 'Mths supply',
                 'Median DOM', 'Absorbed ≤21d', 'For sale'], rows,
           widths=[3.2, 2.1, 2.2, 2.2, 2.0, 2.3, 1.8])

    _section(doc, 'Listing-share movement · agency level')
    for m in ok:
        roll = m['agency_share']['rolling_90d'][:5]
        if not roll:
            continue
        hp = _para(doc, space_before=4, space_after=2)
        _run(hp, m['suburb'], size=9.5, color=GREEN, bold=True)
        maxpct = max((x['pct'] for x in roll), default=1) or 1
        cur = {x['agency']: x for x in m['agency_share']['current_month']}
        _bar_rows(doc, [{
            'label': x['agency'][:34],
            'pct': x['pct'] / maxpct * 100,
            'color': GREEN,
            'value': (f"{x['pct']}%  (mo: "
                      + (f"{cur[x['agency']]['pct']}%" if x['agency'] in cur else '0%')
                      + ")"),
        } for x in roll])

    _section(doc, 'Opportunity register · stale campaigns with price cuts')
    opp = []
    for m in ok:
        for i in m['stale_flags']['items']:
            opp.append([m['suburb'], i['address'], (f"{i['dom']}d", COOL, True),
                        i.get('price_text'), i['price_cuts']])
    opp.sort(key=lambda x: -(int(re.sub(r'\D', '', str(x[2][0])) or 0)))
    if opp:
        _table(doc, ['Suburb', 'Address', 'On market', 'Current ask', 'Cuts'],
               opp[:25], widths=[3.0, 6.2, 2.0, 3.4, 1.4])
    else:
        p = _para(doc)
        _run(p, 'No flagged opportunities across the selected suburbs right now.',
             color=GREY, italic=True)

    _narrative(doc, narratives, 'house_view')
    _notes(doc, [f for m in ok for f in _collect_flags(m)])
    return doc


def build_monthly_deep_dive(metrics, narratives, user_profile):
    doc = _new_report_doc(
        user_profile, 'Monthly Deep Dive',
        'Discount spread, price-band structure and stale-campaign detail.')
    ok = [m for m in metrics if not m.get('error')]

    _section(doc, 'List-to-sale discount spread')
    _table(doc, ['Suburb', 'Median discount', 'Sample', 'Window'],
           [[(m['suburb'], INK, True),
             (_fmt(m['discount'].get('median_pct'), '%'), WARM, True),
             f"{m['discount'].get('sample_size')} sales",
             f"{m['discount'].get('window_days')} days"] for m in ok],
           widths=[4.0, 4.0, 4.5, 4.5])

    _section(doc, 'Price bands')
    for m in ok:
        hp = _para(doc, space_before=4, space_after=2)
        _run(hp, m['suburb'], size=9.5, color=GREEN, bold=True)
        _table(doc, ['Band', 'For sale', 'Median DOM', 'Median discount'],
               [[b['band'], b['active_count'],
                 _fmt(b.get('dom_median'), 'd'),
                 _fmt(b.get('discount_median_pct'), '%')]
                for b in m['price_bands']],
               widths=[4.5, 3.5, 4.5, 4.5])

    _section(doc, 'Stale campaigns · detail')
    any_items = False
    for m in ok:
        sf = m['stale_flags']
        if not sf['items']:
            continue
        any_items = True
        hp = _para(doc, space_before=4, space_after=2)
        _run(hp, f"{m['suburb']} · over {sf['threshold_days']} days on market",
             size=9.5, color=GREEN, bold=True)
        _table(doc, ['Address', 'On market', 'Current ask', 'Price cuts'],
               [[i['address'], (f"{i['dom']}d", COOL, True),
                 i.get('price_text'), i['price_cuts']] for i in sf['items']],
               widths=[7.5, 2.5, 4.0, 3.0])
    if not any_items:
        p = _para(doc)
        _run(p, 'No stale flagged campaigns across the selected suburbs.',
             color=GREY, italic=True)

    _narrative(doc, narratives, 'discount')

    _section(doc, 'Rental market')
    p = _para(doc)
    _run(p, 'This section activates with the rentals scrape for your account. '
            'No rental figures are simulated in the meantime.', color=GREY,
         italic=True)
    _notes(doc, [f for m in ok for f in _collect_flags(m)])
    return doc


def build_vendor_benchmark(metrics, narratives, user_profile):
    """Vendor-safe: aggregated market facts only — no competitor agencies
    or addresses anywhere."""
    doc = _new_report_doc(
        user_profile, 'Vendor Market Benchmark',
        'How your market is moving — timeframes, pricing outcomes and '
        'absorption, from tracked listing data.')
    ok = [m for m in metrics if not m.get('error')]

    for m in ok:
        sc = m['momentum'].get('score')
        tcol, tfill, tword = _temp(sc)
        c = m['counts']
        hp = _para(doc, space_before=6, space_after=1)
        _run(hp, m['suburb'], size=15, color=GREEN, bold=True)
        _kpi_cards(doc, [
            {'value': tword, 'label': 'Market temperature',
             'color': tcol, 'fill': tfill},
            {'value': str(c['active']), 'label': 'Homes for sale',
             'color': INK, 'fill': FILL_CARD},
            {'value': _fmt(m['velocity'].get('dom_median_active'), 'd'),
             'label': 'Median days on market', 'color': INK, 'fill': FILL_CARD},
            {'value': (_fmt(m['discount'].get('median_pct'), '%')
                       if m['discount'].get('available') else 'n/a'),
             'label': 'Typical list-to-sale', 'color': INK, 'fill': FILL_CARD},
        ])
        _section(doc, 'By price band')
        _table(doc, ['Band', 'Homes for sale', 'Median days on market',
                     'Median list-to-sale'],
               [[b['band'], b['active_count'],
                 _fmt(b.get('dom_median'), ' days'),
                 _fmt(b.get('discount_median_pct'), '%')]
                for b in m['price_bands']],
               widths=[3.6, 4.0, 5.0, 4.4])
        obs = _suburb_observations(m)
        # Vendor-safe: drop the competitor/motivated-vendor line.
        obs = [o for o in obs if 'motivated' not in o and 'direct approach' not in o]
        _observations(doc, obs[:4])

    _narrative(doc, narratives, 'market_conditions')
    _notes(doc, [f for m in ok for f in _collect_flags(m)])
    return doc


REPORT_BUILDERS = {
    'suburb_intelligence': (build_suburb_intelligence, 'Suburb_Intelligence'),
    'director_dashboard': (build_director_dashboard, 'Director_Dashboard'),
    'monthly_deep_dive': (build_monthly_deep_dive, 'Monthly_Deep_Dive'),
    'vendor_benchmark': (build_vendor_benchmark, 'Vendor_Benchmark'),
}


def build_report(report_type, metrics, narratives, user_profile):
    builder, stem = REPORT_BUILDERS[report_type]
    doc = builder(metrics, narratives, user_profile)
    suburb_bit = ''
    if len(metrics) == 1:
        safe = re.sub(r'[^\w\s-]', '', metrics[0].get('suburb') or '')
        suburb_bit = '_' + safe.strip().replace(' ', '_') if safe else ''
    stamp = datetime.utcnow().strftime('%Y%m%d')
    return doc, f"SuburbDesk_{stem}{suburb_bit}_{stamp}.docx"
