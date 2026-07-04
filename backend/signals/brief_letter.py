"""SENTINEL S4 — per-signal letter builder for brief items.

Renders a .docx letter for one vendor_signal, with the body ADAPTED to the
signal type (a withdrawn letter is not a street-record letter — handoff
requirement). Uses python-docx (existing dependency) and the calling
agent's real profile; if the profile is incomplete we fall back to
neutral phrasing rather than placeholder contact details (UX-1 rule:
never print fake phone/email on a client letter).
"""
import json
import logging

from docx import Document
from docx.shared import Pt

from database import get_db

logger = logging.getLogger(__name__)


def _letter_flavour(reasons):
    """Pick the body template from the signal's reason codes."""
    text = ' '.join(reasons).lower()
    if 'withdrawn' in text:
        return 'withdrawn'
    if 'sales in the street' in text:
        return 'street_momentum'
    if 'price drop' in text:
        return 'price_drops'
    if 'holding' in text:
        return 'long_hold'
    return 'generic'


_BODIES = {
    'withdrawn': (
        "I noticed that your property at {address} was recently taken off "
        "the market without selling. That can be a frustrating experience, "
        "and it often comes down to strategy rather than the home itself.",
        "If you are still considering a sale — now or later this year — I "
        "would welcome the chance to share what I believe a fresh approach "
        "could achieve for {address}.",
    ),
    'street_momentum': (
        "There has been notable sales activity on your street recently, "
        "and buyer attention in {suburb} is currently focused on homes "
        "like {address}.",
        "If you have ever wondered what that activity means for the value "
        "of your own home, I would be glad to prepare a no-obligation "
        "appraisal while the street is in buyers' minds.",
    ),
    'price_drops': (
        "I have been following the campaign at {address} and noticed the "
        "asking price has been adjusted more than once.",
        "If the current approach is not producing the result you hoped "
        "for, I would welcome a conversation about what a different "
        "strategy could look like — no obligation, of course.",
    ),
    'long_hold': (
        "You have owned {address} for a good number of years, and the "
        "{suburb} market has moved considerably in that time.",
        "Many owners in your position are surprised by what their home "
        "is worth today. I would be glad to prepare a confidential, "
        "no-obligation appraisal whenever it suits you.",
    ),
    'generic': (
        "I work with home owners in {suburb} and wanted to introduce "
        "myself in relation to your property at {address}.",
        "If a sale is something you may consider — now or in the future — "
        "I would be glad to share a view on the current market, with no "
        "obligation.",
    ),
}


def build_brief_letter(signal_id, user):
    """Render the .docx for one signal. Returns (Document, filename) or
    (None, None) when the signal doesn't exist."""
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT v.id, v.address, v.reason_codes, s.name AS suburb "
            "FROM vendor_signals v LEFT JOIN suburbs s ON s.id = v.suburb_id "
            "WHERE v.id = ?", (signal_id,)
        ).fetchone()
        if not row:
            return None, None
        sig = dict(row)
    finally:
        conn.close()

    try:
        reasons = json.loads(sig.get('reason_codes') or '[]')
    except Exception:
        reasons = []
    flavour = _letter_flavour(reasons)
    address = sig['address'] or 'your property'
    suburb = sig['suburb'] or 'your area'

    agent_name = ' '.join(p for p in [
        (user.get('first_name') or '').strip(),
        (user.get('last_name') or '').strip()] if p).strip()
    agent_phone = (user.get('phone') or '').strip()
    agent_email = (user.get('email') or '').strip()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)

    doc.add_paragraph('Dear Home Owner,')
    doc.add_paragraph('')
    for para in _BODIES[flavour]:
        doc.add_paragraph(para.format(address=address, suburb=suburb))
        doc.add_paragraph('')
    doc.add_paragraph('Kind regards,')
    doc.add_paragraph('')
    # UX-1: only real contact details, never placeholders. Missing fields
    # are simply omitted.
    if agent_name:
        p = doc.add_paragraph()
        p.add_run(agent_name).bold = True
    for line in (agent_phone, agent_email):
        if line:
            doc.add_paragraph(line)

    safe = ''.join(c for c in address if c.isalnum() or c in (' ', '-'))
    safe = safe.strip().replace(' ', '_')[:60] or f"signal_{sig['id']}"
    return doc, f"brief_{flavour}_{safe}.docx"
